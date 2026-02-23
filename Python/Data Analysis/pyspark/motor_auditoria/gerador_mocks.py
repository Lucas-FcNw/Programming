"""
COMPONENTE A — Gerador de Mocks Realistas
==========================================
Cria ambiente artificial porém realista para testes do motor de auditoria.
Gera:
  - DOCX com pseudocódigos variados (linguagem humana imperfeita)
  - Notebook com código PySpark (com divergências propositais)
"""

import os
import sys

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import nbformat as nbf
except ImportError as e:
    print(f"ERRO: Falta instalar bibliotecas. Detalhe: {e}")
    print("Execute: pip install python-docx nbformat")
    sys.exit(1)


# ═══════════════════════════════════════════════════════════════
#  DEFINIÇÃO DAS POLÍTICAS (Fonte da verdade para geração)
# ═══════════════════════════════════════════════════════════════

POLITICAS_DOC = [
    # Política 1 — linguagem direta
    "SE score maior que 800, classificar cliente como SUPER_VIP. Caso contrário, classificar como COMUM.",

    # Política 2 — linguagem invertida
    "SE a renda do cliente for menor que 1000 reais, o crédito deve ser NEGADO. Senão, encaminhar para ANALISE.",

    # Política 3 — linguagem informal, sem operador explícito
    "Clientes com idade igual a 18 anos recebem um bônus de 50 pontos no score final.",

    # Política 4 — referência a estado
    "SE o estado do cliente é SP, aplicar taxa de 2%. Para outros estados, aplicar taxa de 10%.",

    # Política 5 — operador IN com lista
    "SE o código do produto estiver entre 10, 20 ou 30, a categoria deve ser ESPECIAL. Caso contrário, NORMAL.",

    # Política 6 — condição composta E
    "SE score acima de 700 E renda acima de 2000, o cliente recebe status PREMIUM.",

    # Política 7 — condição composta OU
    "SE idade menor que 21 OU idade maior que 65, aplicar fator de risco 1.5. Senão, fator 1.0.",

    # Política 8 — política que NÃO será implementada (ausente no notebook)
    "SE o cliente possui mais de 3 contratos ativos, bloquear novas contratações.",

    # Política 9 — linguagem bem diferente, operador implícito
    "Clientes do Rio de Janeiro devem pagar uma sobretaxa de 5% sobre o valor aprovado.",

    # Política 10 — condição com string e fallback
    "SE o tipo de contrato for LEASING, o prazo máximo é 48 meses. Para outros tipos, 60 meses.",
]

# Código PySpark com DIVERGÊNCIAS PROPOSITAIS
CELULAS_NOTEBOOK = [
    # Célula 0 — Setup (não é política)
    {
        "tipo": "setup",
        "codigo": (
            "# Setup do ambiente\n"
            "from pyspark.sql import functions as F\n"
            "from pyspark.sql.types import *\n"
            "import datetime\n"
        ),
    },
    # Célula 1 — Política 1: CORRETA
    {
        "tipo": "politica",
        "codigo": (
            "# Classificação de clientes por score\n"
            "df_result = df_input.withColumn(\n"
            "    'cat',\n"
            "    F.when(F.col('score') > 800, 'SUPER_VIP')\n"
            "     .otherwise('COMUM')\n"
            ")\n"
        ),
    },
    # Célula 2 — Política 2: DIVERGENTE (< 1001 em vez de < 1000)
    {
        "tipo": "politica",
        "codigo": (
            "# Regra de concessão por renda\n"
            "df_credito = df_input.withColumn(\n"
            "    'status',\n"
            "    F.when(F.col('renda') < 1001, 'NEGADO')\n"
            "     .otherwise('ANALISE')\n"
            ")\n"
        ),
    },
    # Célula 3 — Política 3: DIVERGENTE (bônus 10 em vez de 50)
    {
        "tipo": "politica",
        "codigo": (
            "# Ajuste de score por idade\n"
            "df_bonus = df_input.withColumn(\n"
            "    'score_final',\n"
            "    F.when(F.col('idade') == 18, F.col('score') + 10)\n"
            "     .otherwise(F.col('score'))\n"
            ")\n"
        ),
    },
    # Célula 4 — Política 4: DIVERGENTE (estado RJ em vez de SP, taxa 5% em vez de 2%)
    {
        "tipo": "politica",
        "codigo": (
            "# Taxa por região\n"
            "df_taxa = df_input.withColumn(\n"
            "    'taxa',\n"
            "    F.when(F.col('estado') == 'RJ', 0.05)\n"
            "     .otherwise(0.1)\n"
            ")\n"
        ),
    },
    # Célula 5 — Política 5: CORRETA
    {
        "tipo": "politica",
        "codigo": (
            "# Categorização por código de produto\n"
            "df_cat = df_input.withColumn(\n"
            "    'categoria',\n"
            "    F.when(F.col('codigo').isin(10, 20, 30), 'ESPECIAL')\n"
            "     .otherwise('NORMAL')\n"
            ")\n"
        ),
    },
    # Célula 6 — Política 6: CORRETA
    {
        "tipo": "politica",
        "codigo": (
            "# Status premium por score e renda\n"
            "df_premium = df_input.withColumn(\n"
            "    'status_cli',\n"
            "    F.when(\n"
            "        (F.col('score') > 700) & (F.col('renda') > 2000),\n"
            "        'PREMIUM'\n"
            "    ).otherwise('NORMAL')\n"
            ")\n"
        ),
    },
    # Célula 7 — Política 7: DIVERGENTE (operador E em vez de OU)
    {
        "tipo": "politica",
        "codigo": (
            "# Fator de risco por faixa etária\n"
            "df_risco = df_input.withColumn(\n"
            "    'fator_risco',\n"
            "    F.when(\n"
            "        (F.col('idade') < 21) & (F.col('idade') > 65),\n"
            "        1.5\n"
            "    ).otherwise(1.0)\n"
            ")\n"
        ),
    },
    # Célula 8 — Código morto (não é política, é auxiliar)
    {
        "tipo": "auxiliar",
        "codigo": (
            "# Funções auxiliares de log\n"
            "def log_execution(step_name):\n"
            "    print(f'[LOG] Executando: {step_name}')\n"
            "    return True\n"
            "\n"
            "log_execution('pipeline_credito')\n"
        ),
    },
    # Célula 9 — Política 9: PARCIAL (sobretaxa de 3% em vez de 5%)
    {
        "tipo": "politica",
        "codigo": (
            "# Sobretaxa regional\n"
            "df_sobretaxa = df_input.withColumn(\n"
            "    'sobretaxa',\n"
            "    F.when(F.col('estado') == 'RJ', 0.03)\n"
            "     .otherwise(0.0)\n"
            ")\n"
        ),
    },
    # Célula 10 — Política 10: CORRETA
    {
        "tipo": "politica",
        "codigo": (
            "# Prazo máximo por tipo de contrato\n"
            "df_prazo = df_input.withColumn(\n"
            "    'prazo_max',\n"
            "    F.when(F.col('tipo_contrato') == 'LEASING', 48)\n"
            "     .otherwise(60)\n"
            ")\n"
        ),
    },
    # Célula 11 — Código incompleto (só filtro, sem ação)
    {
        "tipo": "incompleto",
        "codigo": (
            "# Tentativa de filtro (incompleto)\n"
            "df_filtrado = df_input.filter(F.col('score') > 500)\n"
            "# TODO: aplicar regra de segmentação\n"
        ),
    },
]


def gerar_docx(caminho_saida: str):
    """Gera documento Word com pseudocódigos de políticas."""
    doc = docx.Document()

    titulo = doc.add_heading("Políticas de Crédito — Especificação", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Este documento descreve as regras de negócio para o motor de crédito. "
        "As políticas devem ser implementadas no pipeline PySpark."
    )
    doc.add_paragraph("")

    for i, texto in enumerate(POLITICAS_DOC, 1):
        p = doc.add_paragraph()
        run_titulo = p.add_run(f"Política {i}: ")
        run_titulo.bold = True
        p.add_run(texto)

    doc.add_paragraph("")
    doc.add_paragraph("— Fim do documento de especificação —")

    doc.save(caminho_saida)
    print(f"  ✅ DOCX criado: {caminho_saida}")


def gerar_notebook(caminho_saida: str):
    """Gera notebook Jupyter com implementações PySpark (com divergências)."""
    nb = nbf.v4.new_notebook()
    nb.metadata = {
        "kernelspec": {
            "display_name": "Python 3",
            "language": "python",
            "name": "python3",
        }
    }

    # Célula markdown inicial
    nb.cells.append(
        nbf.v4.new_markdown_cell(
            "# Pipeline de Crédito\n"
            "Implementação das regras de negócio para análise de crédito."
        )
    )

    for item in CELULAS_NOTEBOOK:
        nb.cells.append(nbf.v4.new_code_cell(item["codigo"]))

    with open(caminho_saida, "w", encoding="utf-8") as f:
        nbf.write(nb, f)

    print(f"  ✅ Notebook criado: {caminho_saida}")


def gerar_tudo(diretorio: str | None = None):
    """Gera todos os arquivos mock no diretório especificado."""
    if diretorio is None:
        diretorio = os.path.dirname(os.path.abspath(__file__))

    os.makedirs(diretorio, exist_ok=True)

    print("=" * 50)
    print("  GERADOR DE MOCKS — Motor de Auditoria")
    print("=" * 50)
    print(f"  Diretório: {diretorio}\n")

    caminho_docx = os.path.join(diretorio, "regras_mock.docx")
    caminho_nb = os.path.join(diretorio, "notebook_mock.ipynb")

    gerar_docx(caminho_docx)
    gerar_notebook(caminho_nb)

    print(f"\n  📊 Políticas no DOC:      {len(POLITICAS_DOC)}")
    print(f"  📓 Células no Notebook:   {len(CELULAS_NOTEBOOK)}")
    print(f"  ⚠️  Divergências plantadas: 4 (P2, P3, P4, P7)")
    print(f"  ❌ Política ausente:       1 (P8)")
    print(f"  🗑️  Código morto:           2 (auxiliar + incompleto)")
    print("=" * 50)

    return caminho_docx, caminho_nb


if __name__ == "__main__":
    gerar_tudo()
